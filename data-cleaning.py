import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# Function to check if run contains red text
def is_red(run):
    if run.font.color and run.font.color.rgb:
        return run.font.color.rgb == (255, 0, 0)  # RGB value for red

# Function to remove images, red text, and empty lines
def clean_document(doc):
    for para in doc.paragraphs:
        new_runs = []
        for run in para.runs:
            if is_red(run):
                new_runs.append(run.text)

        para.clear()  # Clear the paragraph's text
        for text in new_runs:
            para.add_run(text)  # Re-add the non-red text

    # Remove images
    for rel in list(doc.part.rels.values()):
        if "image" in rel.target_ref:
            del doc.part.rels[rel.rId]

    # Remove empty paragraphs
    for para in doc.paragraphs:
        if not para.text.strip():
            p = para._element
            p.getparent().remove(p)

    return doc

# Function to convert and save the document as txt file
def save_as_txt(doc_path, txt_path):
    doc = Document(doc_path)
    with open(txt_path, 'w', encoding='utf-8') as txt_file:
        for para in doc.paragraphs:
            txt_file.write(para.text + '\n')

# Create 'articles' directory if it doesn't exist
output_dir = "articles"
os.makedirs(output_dir, exist_ok=True)

# Paths to the documents
doc_paths = [
    "input-files/Налоговый кодекс РК.docx"
    # "input-files/О банках и банковской деятельности в Республике Казахстан.docx",
    # "input-files/О персональных данных и их защите.docx",
    # "input-files/ОСвязи.docx"
]

# Clean and save each document
for doc_path in doc_paths:
    doc = Document(doc_path)
    cleaned_doc = clean_document(doc)
    
    # Save the cleaned docx in 'articles' folder
    cleaned_doc_path = os.path.join(output_dir, os.path.basename(doc_path).replace(".docx", "_Cleaned.docx"))
    cleaned_doc.save(cleaned_doc_path)
    
    # Convert to txt and save in 'articles' folder
    # txt_file_name = os.path.basename(cleaned_doc_path).replace(".docx", ".txt")

    # txt_path = os.path.join(output_dir, cleaned_doc)
    # save_as_txt(cleaned_doc_path, txt_path)

# List the paths of the saved text files
txt_files = [os.path.join(output_dir, f) for f in os.listdir(output_dir) if f.endswith('.docx')]
txt_files
